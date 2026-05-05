"""
docx_generator.py — Layer 3 Tool: Cover Letter → DOCX

Converts plain text cover letter into a clean, professionally formatted .docx file.
Uses python-docx. No LLM involved.
"""

import os
import re
from pathlib import Path
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ─── Naming Helper ────────────────────────────────────────────────────────────

def _slugify(text: str) -> str:
    """Convert text to a filename-safe slug."""
    text = text.lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_-]+", "_", text)
    return text[:40]  # Cap length


def _build_output_path(job_title: str, company: str, output_dir: Path) -> Path:
    """
    Build the output path for the cover letter DOCX.
    Format: output/{company}_{job_title}_{YYYY-MM-DD}/cover_letter.docx
    """
    date_str = datetime.today().strftime("%Y-%m-%d")
    folder_name = f"{_slugify(company)}_{_slugify(job_title)}_{date_str}"
    folder_path = output_dir / folder_name

    # Handle duplicate folder names
    if folder_path.exists():
        counter = 2
        while folder_path.exists():
            folder_path = output_dir / f"{folder_name}_{counter}"
            counter += 1

    folder_path.mkdir(parents=True, exist_ok=True)
    return folder_path / "cover_letter.docx"


# ─── DOCX Generation ─────────────────────────────────────────────────────────

def generate_docx(
    cover_letter_text: str,
    job_title: str,
    company: str,
    output_dir: Optional[Path] = None,
) -> Path:
    """
    Generate a professional cover letter .docx from plain text.

    Args:
        cover_letter_text: Plain text cover letter (paragraphs separated by newlines).
        job_title: Job title (used in filename + header).
        company: Company name (used in filename + header).
        output_dir: Directory to save to. Defaults to ./output.

    Returns:
        Path to the generated .docx file.
    """
    if output_dir is None:
        output_dir = Path(__file__).parent.parent / "output"
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    output_path = _build_output_path(job_title, company, output_dir)

    doc = Document()

    # ── Page Margins ────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    # ── Default Style ────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # ── Date ────────────────────────────────────────────────────
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date_para.add_run(datetime.today().strftime("%B %d, %Y"))
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)

    doc.add_paragraph()  # Spacer

    # ── Body: Split on double newlines (paragraphs) ───────────────
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", cover_letter_text.strip()) if p.strip()]

    for i, para_text in enumerate(paragraphs):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = para.add_run(para_text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)

        # Add spacing between paragraphs
        para.paragraph_format.space_after = Pt(10)

        # First paragraph: slightly larger indent for elegance
        if i == 0:
            para.paragraph_format.space_before = Pt(6)

    doc.save(str(output_path))
    return output_path


# ─── CLI Test ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys
    sample_text = """Dear Hiring Manager,

I am writing to express my strong interest in the Senior Software Engineer position at Acme Corp. With over five years of experience building scalable Python microservices and leading cross-functional teams, I am confident in my ability to contribute meaningfully from day one.

Throughout my career at TechStartup, I delivered production-grade REST APIs serving over 50,000 daily active users, leveraging Python, Docker, and AWS to ensure reliability and performance. I have a track record of shipping features on time while maintaining high code quality standards.

I am particularly excited about Acme Corp's mission to democratize access to financial data. I believe my background in fintech and distributed systems makes me an excellent fit for this role.

Thank you for considering my application. I look forward to the opportunity to discuss how I can contribute to your team.

Sincerely,
Your Name"""

    output = generate_docx(
        cover_letter_text=sample_text,
        job_title="Senior Software Engineer",
        company="Acme Corp",
    )
    print(f"✅ Cover letter generated: {output}")
